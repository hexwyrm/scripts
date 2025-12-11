# Cybersecurity Digest Generator - by hexwyrm
# ---------------------------------------------------

import os
import requests
from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from docx.shared import Pt

# ---------------------------------------------------
# RSS Source Definitions
RSS_SOURCES = {
    "The Hacker News": "https://feeds.feedburner.com/TheHackersNews",
    "Krebs on Security": "https://krebsonsecurity.com/feed/",
    "BleepingComputer": "https://www.bleepingcomputer.com/feed/",
    "HackRead": "https://hackread.com/feed/",
    "Dark Reading": "https://www.darkreading.com/rss.xml",
    "CSO Online": "https://www.csoonline.com/index.rss"
}

# ---------------------------------------------------
# Function: Parse RSS Feed
def parse_rss(source_name, url, max_articles=5):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, "xml")
        items = soup.find_all("item")[:max_articles]

        articles = []
        for item in items:
            title = item.title.text.strip()
            link = item.link.text.strip()
            description = item.description.text.strip()
            articles.append({
                "title": title,
                "url": link,
                "summary": description
            })
        return source_name, articles
    except Exception as e:
        print(f"[!] Failed to parse {source_name}: {e}")
        return source_name, []

# ---------------------------------------------------
# Function: Save to DOCX
def save_to_docx(news_dict, filename = f"cybersecurity_digest_{datetime.today().strftime('%Y%m%d')}.docx"):
    script_dir = os.path.dirname(os.path.abspath(__file__))
    full_path = os.path.join(script_dir, filename)

    doc = Document()
    doc.add_heading(f"🛡️ Cybersecurity Digest – {datetime.today().strftime('%B %d, %Y')}", level=1)

    for source, articles in news_dict.items():
        doc.add_heading(source, level=2)
        for article in articles:
            doc.add_heading(article['title'], level=3)
            doc.add_paragraph(f"📎 Read more: {article['url']}", style='List Bullet')
            doc.add_paragraph(f"📝 {article['summary']}")

    doc.save(full_path)

# ---------------------------------------------------
# Main Execution
def main():
    news_dict = {}
    max_articles_per_source = 5

    for source, url in RSS_SOURCES.items():
        name, articles = parse_rss(source, url, max_articles=max_articles_per_source)
        news_dict[name] = articles

    save_to_docx(news_dict)
    total_articles = sum(len(v) for v in news_dict.values())
    print(f"[✓] Saved {total_articles} articles to cybersecurity_digest.docx")

# ---------------------------------------------------
# Run the script
if __name__ == "__main__":
    main()
